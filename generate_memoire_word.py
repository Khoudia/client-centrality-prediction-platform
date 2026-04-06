# -*- coding: utf-8 -*-
"""
Générateur Word du mémoire DU SDA 2025-2026
Université Paris-Sorbonne - Hôtel Aurore Paris Gare de Lyon
Analyse Réseau & Satisfaction Client

Usage:
    python generate_memoire_word.py
"""

import re
import os
import sys
from pathlib import Path

# Force UTF-8 output
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────────────────
# CHEMINS
# ──────────────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
MD_FILE  = BASE_DIR / "docs" / "MEMOIRE_MASTER_DATA_SCIENCE_2026.md"
OUT_FILE = BASE_DIR / "docs" / "MEMOIRE_MASTER_DATA_SCIENCE_2026.docx"

# ──────────────────────────────────────────────────────────────────────────────
# PALETTE COULEURS
# ──────────────────────────────────────────────────────────────────────────────
BLEU_SORBONNE  = RGBColor(0x1F, 0x39, 0x64)  # Bleu marine Sorbonne
BLEU_MOYEN     = RGBColor(0x2E, 0x74, 0xB5)  # Bleu Microsoft Word
BLEU_CLAIR_HEX = "DEEAF1"                     # En-tête tableau
BLEU_ALT_HEX   = "EBF3FB"                     # Lignes alternées
HEADER_HEX     = "1F3964"                     # En-tête tableau sombre
GRIS_TEXTE     = RGBColor(0x50, 0x50, 0x50)
GRIS_CODE_HEX  = "F2F2F2"
BLANC          = RGBColor(0xFF, 0xFF, 0xFF)


# ──────────────────────────────────────────────────────────────────────────────
# HELPERS XML
# ──────────────────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str):
    """Applique une couleur de fond à une cellule."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    # Supprimer ancien shading si présent
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shd)


def _set_para_shading(para, fill_hex: str):
    """Applique une couleur de fond à un paragraphe."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    pPr.append(shd)


def _add_field(run, field_code: str):
    """Insère un champ Word (PAGE, NUMPAGES…) dans un run."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_hline(doc: Document, color_hex: str = "2E74B5", thickness: int = 8):
    """Ajoute un filet horizontal décoratif."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ──────────────────────────────────────────────────────────────────────────────
# CONFIGURATION DES STYLES
# ──────────────────────────────────────────────────────────────────────────────

def setup_styles(doc: Document):
    """Définit tous les styles personnalisés du document."""
    styles = doc.styles

    # ── Normal ──
    normal = styles["Normal"]
    normal.font.name            = "Calibri"
    normal.font.size            = Pt(11)
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = Pt(14)

    # ── Heading 1 ──
    h1 = styles["Heading 1"]
    h1.font.name      = "Calibri Light"
    h1.font.size      = Pt(20)
    h1.font.bold      = True
    h1.font.color.rgb = BLEU_SORBONNE
    h1.paragraph_format.space_before      = Pt(24)
    h1.paragraph_format.space_after       = Pt(12)
    h1.paragraph_format.keep_with_next    = True
    h1.paragraph_format.page_break_before = True

    # ── Heading 2 ──
    h2 = styles["Heading 2"]
    h2.font.name      = "Calibri Light"
    h2.font.size      = Pt(15)
    h2.font.bold      = True
    h2.font.color.rgb = BLEU_MOYEN
    h2.paragraph_format.space_before   = Pt(18)
    h2.paragraph_format.space_after    = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # ── Heading 3 ──
    h3 = styles["Heading 3"]
    h3.font.name      = "Calibri"
    h3.font.size      = Pt(13)
    h3.font.bold      = True
    h3.font.color.rgb = BLEU_MOYEN
    h3.paragraph_format.space_before   = Pt(14)
    h3.paragraph_format.space_after    = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # ── Heading 4 ──
    h4 = styles["Heading 4"]
    h4.font.name      = "Calibri"
    h4.font.size      = Pt(11)
    h4.font.bold      = True
    h4.font.italic    = True
    h4.font.color.rgb = GRIS_TEXTE
    h4.paragraph_format.space_before   = Pt(10)
    h4.paragraph_format.space_after    = Pt(4)
    h4.paragraph_format.keep_with_next = True

    # ── Code (personnalisé) ──
    if "CodeBlock" not in [s.name for s in styles]:
        cs = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cs = styles["CodeBlock"]
    cs.base_style          = styles["Normal"]
    cs.font.name           = "Courier New"
    cs.font.size           = Pt(8.5)
    cs.paragraph_format.space_before  = Pt(6)
    cs.paragraph_format.space_after   = Pt(6)
    cs.paragraph_format.left_indent   = Cm(0.6)
    cs.paragraph_format.right_indent  = Cm(0.3)
    cs.paragraph_format.line_spacing  = Pt(12)

    # ── Quote ──
    if "BlockQuote" not in [s.name for s in styles]:
        qs = styles.add_style("BlockQuote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        qs = styles["BlockQuote"]
    qs.base_style = styles["Normal"]
    qs.font.name  = "Calibri"
    qs.font.size  = Pt(10.5)
    qs.font.italic = True
    qs.font.color.rgb = GRIS_TEXTE
    qs.paragraph_format.left_indent   = Cm(1.8)
    qs.paragraph_format.right_indent  = Cm(1.0)
    qs.paragraph_format.space_before  = Pt(8)
    qs.paragraph_format.space_after   = Pt(8)


# ──────────────────────────────────────────────────────────────────────────────
# PAGE DE TITRE
# ──────────────────────────────────────────────────────────────────────────────

def build_title_page(doc: Document):
    """Crée la page de titre académique."""

    def _centre_run(text, size, bold=False, italic=False, color=None, font="Calibri Light"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
        return p

    # Espace supérieur
    for _ in range(2):
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # Logo textuel université
    _centre_run("✦  UNIVERSITÉ PARIS-SORBONNE  ✦", 14, bold=True, color=BLEU_SORBONNE)
    _centre_run("DU SDA — Promotion 2025-2026", 12, color=BLEU_MOYEN)

    _add_hline(doc, color_hex="1F3964", thickness=12)

    for _ in range(1):
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # Titre principal
    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titre.paragraph_format.space_before = Pt(10)
    p_titre.paragraph_format.space_after  = Pt(6)
    r = p_titre.add_run("ANALYSE RÉSEAU ET MODÉLISATION\nDE LA SATISFACTION CLIENT\nDANS LE SECTEUR HÔTELIER")
    r.font.name      = "Calibri Light"
    r.font.size      = Pt(24)
    r.font.bold      = True
    r.font.color.rgb = BLEU_SORBONNE

    # Sous-titre
    _centre_run("Application à l'Hôtel Aurore Paris Gare de Lyon",
                14, italic=True, color=BLEU_MOYEN)

    _add_hline(doc, color_hex="2E74B5", thickness=8)

    for _ in range(1):
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(10)

    # Tableau des informations
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    infos = [
        ("Auteur",               "[Votre Nom Prénom]"),
        ("Directeur de mémoire", "[Nom du Directeur]"),
        ("Date de soutenance",   "Juin 2026"),
        ("UFR",                  "Sciences Mathématiques et Informatique"),
        ("Spécialité",           "Data Science & Intelligence Artificielle"),
    ]
    for i, (label, val) in enumerate(infos):
        row = tbl.rows[i]
        # Colonne label
        c0 = row.cells[0]
        c0.text = ""
        _set_cell_shading(c0, BLEU_CLAIR_HEX)
        rr = c0.paragraphs[0].add_run(label)
        rr.font.name = "Calibri"; rr.font.size = Pt(10); rr.font.bold = True
        rr.font.color.rgb = BLEU_SORBONNE
        c0.width = Cm(5)
        # Colonne valeur
        c1 = row.cells[1]
        c1.text = ""
        rr2 = c1.paragraphs[0].add_run(val)
        rr2.font.name = "Calibri"; rr2.font.size = Pt(10)
        c1.width = Cm(10)

    for _ in range(2):
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # Citation épigraphe
    p_q = doc.add_paragraph()
    p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_q.paragraph_format.left_indent  = Cm(2.5)
    p_q.paragraph_format.right_indent = Cm(2.5)
    p_q.paragraph_format.space_before = Pt(10)
    r_q = p_q.add_run(
        "\u00ab\u00a0La valeur d\u2019un individu dans un r\u00e9seau ne r\u00e9side pas seulement dans ce qu\u2019il est,\n"
        "mais dans la fa\u00e7on dont il est connect\u00e9 aux autres.\u00a0\u00bb\n"
        "\u2014 Mark Newman, Networks: An Introduction, 2010"
    )
    r_q.font.name      = "Calibri"
    r_q.font.size      = Pt(10)
    r_q.font.italic    = True
    r_q.font.color.rgb = GRIS_TEXTE

    doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
# FORMATAGE INLINE MARKDOWN
# ──────────────────────────────────────────────────────────────────────────────

_INLINE_PATTERN = re.compile(
    r"(\*\*\*[^*]+?\*\*\*"      # ***gras+italique***
    r"|\*\*[^*]+?\*\*"           # **gras**
    r"|\*[^*]+?\*"               # *italique*
    r"|`[^`]+?`"                 # `code`
    r"|\[([^\]]+)\]\([^\)]+\)"  # [texte](url)
    r")"
)


def _write_inline(para, text: str, base_bold=False, base_italic=False,
                  base_size=None, base_color=None, base_font="Calibri"):
    """Écrit du texte avec formatage inline Markdown dans un paragraphe."""
    parts = _INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        r = para.add_run()
        r.font.name = base_font
        if base_size:
            r.font.size = Pt(base_size)
        if base_color:
            r.font.color.rgb = base_color
        if base_bold:
            r.bold = True
        if base_italic:
            r.italic = True

        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            r.text = part[3:-3]; r.bold = True; r.italic = True
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            r.text = part[2:-2]; r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r.text = part[1:-1]; r.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r.text = part[1:-1]
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        elif re.match(r"\[([^\]]+)\]\([^\)]+\)", part):
            m = re.match(r"\[([^\]]+)\]\([^\)]+\)", part)
            r.text = m.group(1)
            r.underline = True
            r.font.color.rgb = BLEU_MOYEN
        else:
            r.text = part


def _plain(text: str) -> str:
    """Retire tous les marqueurs Markdown d'une chaîne."""
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text.strip()


# ──────────────────────────────────────────────────────────────────────────────
# RENDU TABLEAUX MARKDOWN
# ──────────────────────────────────────────────────────────────────────────────

def _render_table(doc: Document, md_lines: list):
    """Convertit un bloc de lignes Markdown en tableau Word stylisé."""
    # Extraire uniquement les lignes de données (pas les séparateurs ---)
    data_rows = []
    for line in md_lines:
        line = line.strip()
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue  # ligne séparateur
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        data_rows.append(cells)

    if not data_rows:
        return

    ncols = max(len(r) for r in data_rows)
    data_rows = [r + [""] * (ncols - len(r)) for r in data_rows]

    tbl = doc.add_table(rows=len(data_rows), cols=ncols)
    tbl.style  = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(data_rows):
        row = tbl.rows[i]
        is_header = (i == 0)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

            if is_header:
                _set_cell_shading(cell, HEADER_HEX)
                _write_inline(p, cell_text,
                              base_bold=True, base_size=9,
                              base_color=BLANC)
            elif i % 2 == 0:
                _set_cell_shading(cell, BLEU_ALT_HEX)
                _write_inline(p, cell_text, base_size=9)
            else:
                _write_inline(p, cell_text, base_size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ──────────────────────────────────────────────────────────────────────────────
# RENDU BLOCS DE CODE
# ──────────────────────────────────────────────────────────────────────────────

def _render_code_block(doc: Document, code_lines: list, lang: str = ""):
    """Rend un bloc de code avec fond gris et police monospace."""
    if not code_lines:
        return

    # Entête langue (si disponible)
    if lang.strip():
        p_lang = doc.add_paragraph()
        p_lang.paragraph_format.space_after  = Pt(0)
        p_lang.paragraph_format.space_before = Pt(6)
        p_lang.paragraph_format.left_indent  = Cm(0.6)
        _set_para_shading(p_lang, "1F3964")
        rl = p_lang.add_run(lang.upper())
        rl.font.name  = "Courier New"
        rl.font.size  = Pt(7)
        rl.font.bold  = True
        rl.font.color.rgb = BLANC

    full_code = "\n".join(code_lines)
    p = doc.add_paragraph(style="CodeBlock")
    _set_para_shading(p, GRIS_CODE_HEX)

    # Bordure gauche bleue
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_bdr = OxmlElement("w:left")
    left_bdr.set(qn("w:val"),   "single")
    left_bdr.set(qn("w:sz"),    "12")
    left_bdr.set(qn("w:space"), "4")
    left_bdr.set(qn("w:color"), "2E74B5")
    pBdr.append(left_bdr)
    pPr.append(pBdr)

    r = p.add_run(full_code)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ──────────────────────────────────────────────────────────────────────────────
# PARSER PRINCIPAL
# ──────────────────────────────────────────────────────────────────────────────

def parse_markdown(doc: Document, md_text: str):
    """
    Parse le Markdown ligne par ligne et génère le document Word.
    Gère : titres, listes, tableaux, blocs de code, citations, formules, texte.
    """
    lines = md_text.splitlines()
    total = len(lines)
    i = 0

    in_code  = False
    code_buf = []
    code_lang = ""

    in_table  = False
    table_buf = []

    # Skip la section "front matter" de titre/résumé déjà rendue en page de titre
    skip_front = True

    while i < total:
        line = lines[i]
        stripped = line.strip()

        # ── 1. Blocs de code ────────────────────────────────────────────────
        fence_m = re.match(r"^(`{3,4}|~{3,4})(.*)", stripped)
        if fence_m:
            if not in_code:
                in_code   = True
                code_lang = fence_m.group(2).strip()
                code_buf  = []
            else:
                # Fermeture
                in_code = False
                if in_table:
                    _render_table(doc, table_buf)
                    table_buf = []; in_table = False
                _render_code_block(doc, code_buf, code_lang)
                code_buf  = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_buf.append(line.rstrip())
            i += 1
            continue

        # ── 2. Flush tableau si la ligne n'est plus un tableau ───────────────
        if in_table and not stripped.startswith("|"):
            _render_table(doc, table_buf)
            table_buf = []; in_table = False
            # Ne pas incrementer i, traiter la ligne courante

        # ── 3. Lignes de tableau ─────────────────────────────────────────────
        if stripped.startswith("|"):
            in_table = True
            table_buf.append(stripped)
            i += 1
            continue

        # ── 4. Lignes vides ──────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── 5. Séparateurs --- (saut de section) ─────────────────────────────
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            _add_hline(doc, color_hex="BFBFBF", thickness=4)
            i += 1
            continue

        # ── 6. Saut de page explicite ─────────────────────────────────────────
        if stripped == "<!-- page_break -->":
            doc.add_page_break()
            i += 1
            continue

        # ── 7. Titres ────────────────────────────────────────────────────────
        m_h = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m_h:
            level = len(m_h.group(1))
            title_raw = m_h.group(2).strip()
            title_text = _plain(title_raw)

            if level == 1:
                # Sauter le tout premier H1 (déjà dans la page de titre)
                if skip_front and "MÉMOIRE" in title_text.upper():
                    i += 1
                    continue
                skip_front = False
                p = doc.add_paragraph(style="Heading 1")
                _write_inline(p, title_raw,
                               base_bold=True, base_size=20,
                               base_color=BLEU_SORBONNE, base_font="Calibri Light")

            elif level == 2:
                # Sauter les sections de front-matter (RÉSUMÉ, ABSTRACT, REMERCIEMENTS, TABLE DES MATIÈRES…)
                FRONT_SECTIONS = {"RÉSUMÉ", "ABSTRACT", "REMERCIEMENTS",
                                  "TABLE DES MATIÈRES", "LISTE DES FIGURES",
                                  "LISTE DES TABLEAUX"}
                skip_front = False
                p = doc.add_paragraph(style="Heading 2")
                _write_inline(p, title_raw,
                               base_bold=True, base_size=15,
                               base_color=BLEU_MOYEN, base_font="Calibri Light")

            elif level == 3:
                skip_front = False
                p = doc.add_paragraph(style="Heading 3")
                _write_inline(p, title_raw,
                               base_bold=True, base_size=13,
                               base_color=BLEU_MOYEN)

            elif level == 4:
                p = doc.add_paragraph(style="Heading 4")
                _write_inline(p, title_raw, base_bold=True, base_italic=True, base_size=11)

            else:
                p = doc.add_paragraph(style="Heading 4")
                _write_inline(p, title_raw, base_bold=True, base_size=10)

            i += 1
            continue

        # ── 8. Citations (>) ─────────────────────────────────────────────────
        if stripped.startswith("> "):
            quote_text = stripped[2:].strip()
            p = doc.add_paragraph(style="BlockQuote")
            _write_inline(p, quote_text, base_italic=True, base_size=10.5,
                          base_color=GRIS_TEXTE)
            i += 1
            continue

        # ── 9. Formules mathématiques $$…$$ ──────────────────────────────────
        if re.match(r"^\$\$", stripped):
            formula = stripped.strip("$").strip()
            if formula:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(6)
                p.paragraph_format.left_indent  = Cm(1)
                r = p.add_run(formula)
                r.font.name      = "Cambria Math"
                r.font.size      = Pt(11)
                r.font.italic    = True
                r.font.color.rgb = BLEU_SORBONNE
            i += 1
            continue

        # ── 10. Listes ────────────────────────────────────────────────────────
        m_li = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if m_li:
            indent = len(m_li.group(1)) // 2
            is_ordered = bool(re.match(r"\d+\.", m_li.group(2)))
            item_text = m_li.group(3)
            style = "List Number" if is_ordered else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent   = Cm(0.6 + indent * 0.6)
            p.paragraph_format.space_after   = Pt(3)
            p.paragraph_format.space_before  = Pt(1)
            _write_inline(p, item_text)
            i += 1
            continue

        # ── 11. Paragraphe texte normal ───────────────────────────────────────
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        _write_inline(p, stripped)
        i += 1

    # Flush tableaux/code restants
    if in_table and table_buf:
        _render_table(doc, table_buf)
    if in_code and code_buf:
        _render_code_block(doc, code_buf, code_lang)


# ──────────────────────────────────────────────────────────────────────────────
# EN-TÊTE ET PIED DE PAGE
# ──────────────────────────────────────────────────────────────────────────────

def setup_header_footer(doc: Document):
    """Configure l'en-tête et le pied de page."""
    section = doc.sections[0]

    # En-tête
    header = section.header
    header.is_linked_to_previous = False
    p_hdr = header.paragraphs[0]
    p_hdr.clear()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_hdr = p_hdr.add_run("Mémoire DU SDA 2025-2026 — Hôtel Aurore Paris")
    r_hdr.font.name      = "Calibri"
    r_hdr.font.size      = Pt(8)
    r_hdr.font.italic    = True
    r_hdr.font.color.rgb = GRIS_TEXTE
    # Filet sous l'en-tête
    pPr = p_hdr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E74B5")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Pied de page
    footer = section.footer
    footer.is_linked_to_previous = False
    p_ftr = footer.paragraphs[0]
    p_ftr.clear()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_left = p_ftr.add_run("Université Paris-Sorbonne  |  ")
    r_left.font.name = "Calibri"; r_left.font.size = Pt(8)
    r_left.font.color.rgb = GRIS_TEXTE

    r_page = p_ftr.add_run("Page ")
    r_page.font.name = "Calibri"; r_page.font.size = Pt(8)
    r_page.font.color.rgb = GRIS_TEXTE
    _add_field(r_page, " PAGE ")

    r_of = p_ftr.add_run(" / ")
    r_of.font.name = "Calibri"; r_of.font.size = Pt(8)
    r_of.font.color.rgb = GRIS_TEXTE
    _add_field(r_of, " NUMPAGES ")

    # Filet au-dessus du pied de page
    pPr2 = p_ftr._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2 = OxmlElement("w:top")
    top2.set(qn("w:val"),   "single")
    top2.set(qn("w:sz"),    "4")
    top2.set(qn("w:space"), "1")
    top2.set(qn("w:color"), "2E74B5")
    pBdr2.append(top2)
    pPr2.append(pBdr2)


# ──────────────────────────────────────────────────────────────────────────────
# MARGES DE PAGE
# ──────────────────────────────────────────────────────────────────────────────

def setup_page(doc: Document):
    """Configure les marges et le format de page."""
    section = doc.sections[0]
    section.page_width   = Cm(21.0)   # A4
    section.page_height  = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


# ──────────────────────────────────────────────────────────────────────────────
# RÉSUMÉ / ABSTRACT
# ──────────────────────────────────────────────────────────────────────────────

def add_resume_page(doc: Document):
    """Ajoute une page de résumé/abstract stylisée."""
    doc.add_page_break()

    def section_box(title, text, lang="FR"):
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(10)
        p_t.paragraph_format.space_after  = Pt(4)
        _set_para_shading(p_t, HEADER_HEX)
        rt = p_t.add_run(f"  {title} ({lang})")
        rt.font.name = "Calibri"; rt.font.size = Pt(12)
        rt.font.bold = True; rt.font.color.rgb = BLANC

        p_body = doc.add_paragraph()
        p_body.paragraph_format.left_indent  = Cm(0.5)
        p_body.paragraph_format.space_before = Pt(4)
        p_body.paragraph_format.space_after  = Pt(8)
        _write_inline(p_body, text, base_size=10)

    RESUME = (
        "Ce mémoire présente une plateforme d\u2019analyse de données hybride conçue pour l\u2019"
        "**Hôtel Aurore Paris Gare de Lyon**, combinant la **théorie des graphes**, la **détection "
        "de communautés** et le **Machine Learning supervisé** pour modéliser et prédire la "
        "satisfaction client. À partir de quatre sources de données réelles (réservations AvailPro, "
        "avis Booking.com et Expedia), un pipeline Python complet a été développé. Le modèle "
        "**Random Forest** atteint **91.3 % d\u2019accuracy** sur la classification high_satisfaction, "
        "avec le PageRank réseau en 3\u1d49 position des variables prédictives. Le réseau de "
        "similarité révèle 5 à 8 communautés de clients interprétables sur le plan métier."
    )

    ABSTRACT = (
        "This thesis presents a hybrid data analytics platform for **Hôtel Aurore Paris Gare de Lyon**, "
        "combining **graph theory**, **community detection**, and **supervised Machine Learning** to "
        "model and predict customer satisfaction. Using four real data sources (AvailPro reservations, "
        "Booking.com and Expedia reviews), a complete Python pipeline was developed. The **Random Forest** "
        "model achieves **91.3% accuracy** on the high_satisfaction classification task, with network "
        "PageRank ranking 3rd among predictive features. The similarity network reveals 5–8 business-"
        "interpretable client communities."
    )

    section_box("RÉSUMÉ", RESUME, "FR")

    p_kw = doc.add_paragraph()
    _set_para_shading(p_kw, BLEU_CLAIR_HEX)
    p_kw.paragraph_format.left_indent = Cm(0.5)
    rkw = p_kw.add_run("Mots-clés : ")
    rkw.font.bold = True; rkw.font.size = Pt(9); rkw.font.name = "Calibri"
    rkw2 = p_kw.add_run(
        "analyse de réseau, théorie des graphes, centralité, communautés, satisfaction client, "
        "Machine Learning, Random Forest, hôtellerie, Python, Streamlit"
    )
    rkw2.font.size = Pt(9); rkw2.font.italic = True; rkw2.font.name = "Calibri"

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    section_box("ABSTRACT", ABSTRACT, "EN")

    p_kw2 = doc.add_paragraph()
    _set_para_shading(p_kw2, BLEU_CLAIR_HEX)
    p_kw2.paragraph_format.left_indent = Cm(0.5)
    rkw3 = p_kw2.add_run("Keywords: ")
    rkw3.font.bold = True; rkw3.font.size = Pt(9); rkw3.font.name = "Calibri"
    rkw4 = p_kw2.add_run(
        "network analysis, graph theory, centrality, customer communities, customer satisfaction, "
        "Machine Learning, Random Forest, hospitality, Python, Streamlit"
    )
    rkw4.font.size = Pt(9); rkw4.font.italic = True; rkw4.font.name = "Calibri"

    doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
# PROGRAMME PRINCIPAL
# ──────────────────────────────────────────────────────────────────────────────

def main():
    print(f"[1/6] Lecture du memoire : {MD_FILE.name}")
    try:
        md_text = MD_FILE.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        md_text = MD_FILE.read_text(encoding="latin-1")
    print(f"      {len(md_text):,} caracteres lus, {md_text.count(chr(10))} lignes")

    print("[2/6] Initialisation du document Word...")
    doc = Document()

    print("[3/6] Configuration (styles, marges, en-tete/pied)...")
    setup_page(doc)
    setup_styles(doc)
    setup_header_footer(doc)

    print("[4/6] Page de titre...")
    build_title_page(doc)

    print("[5/6] Ajout resume/abstract...")
    add_resume_page(doc)

    print("[6/6] Conversion du corps du memoire (Markdown -> Word)...")
    parse_markdown(doc, md_text)

    print(f"      Sauvegarde -> {OUT_FILE.name}")
    doc.save(str(OUT_FILE))

    size_kb = OUT_FILE.stat().st_size / 1024
    # Compter éléments
    n_paras  = len(doc.paragraphs)
    n_tables = len(doc.tables)

    print()
    print("=" * 55)
    print("  CONVERSION TERMINEE AVEC SUCCES")
    print("=" * 55)
    print(f"  Fichier : {OUT_FILE}")
    print(f"  Taille  : {size_kb:.1f} Ko")
    print(f"  Paras   : {n_paras}")
    print(f"  Tableaux: {n_tables}")
    print("=" * 55)
    print()
    print("  Pour ouvrir le fichier :")
    print(f"  Start-Process '{OUT_FILE}'")
    print()


if __name__ == "__main__":
    main()

