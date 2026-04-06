"""
Convertisseur Markdown → Word professionnel
Rapport de stage DU Sorbonne Data Analytics 2025-2026
Hôtel Aurore Paris Gare de Lyon

Génère :
  - RAPPORT_DE_STAGE_SDA_2026.docx  (rapport complet)
  - NOTE_DE_SYNTHESE_SDA_2026.docx  (note de synthèse 2-3 pages)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Chemins ────────────────────────────────────────────────────────────────────
ROOT         = Path(__file__).parent
RAPPORT_MD   = ROOT / "RAPPORT_DE_STAGE_SDA_2026.md"
NOTE_MD      = ROOT / "NOTE_DE_SYNTHESE_SDA_2026.md"
RAPPORT_DOCX = ROOT / "RAPPORT_DE_STAGE_SDA_2026.docx"
NOTE_DOCX    = ROOT / "NOTE_DE_SYNTHESE_SDA_2026.docx"

# ── Palette ────────────────────────────────────────────────────────────────────
COL_BLEU_FONCE = RGBColor(0x1F, 0x3C, 0x88)
COL_BLEU_MOYEN = RGBColor(0x2D, 0x5B, 0xE3)
COL_BLANC      = RGBColor(0xFF, 0xFF, 0xFF)
COL_TEXTE      = RGBColor(0x22, 0x22, 0x22)


# ══════════════════════════════════════════════════════════════════════════════
# Helpers bas niveau
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex6: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6)
    tcPr.append(shd)


def _set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tcB.append(el)
    tcPr.append(tcB)


def _add_page_number(para):
    run = para.add_run()
    for tag, txt in [("begin", None), (None, "PAGE"), ("end", None)]:
        if tag:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        else:
            el = OxmlElement("w:instrText")
            el.text = txt
        run._r.append(el)


def _hrule(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F3C88")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def _run(para, text, bold=False, italic=False, size=11,
         color=None, font="Calibri"):
    r = para.add_run(text)
    r.font.name   = font
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _strip(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text.strip()


def _spacing(para, before=0, after=4):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


# ══════════════════════════════════════════════════════════════════════════════
# Configuration document
# ══════════════════════════════════════════════════════════════════════════════

def _configure(doc: Document):
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    # Pied de page
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, "Hôtel Aurore Paris Gare de Lyon  ·  DU SDA 2025-2026  ·  Page ",
         size=9, color=RGBColor(0x88, 0x88, 0x88))
    _add_page_number(fp)


# ══════════════════════════════════════════════════════════════════════════════
# Page de couverture
# ══════════════════════════════════════════════════════════════════════════════

def _cover(doc: Document, is_note: bool):
    def _ctr(txt, size, bold=False, italic=False, col=None, spa=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, txt, bold=bold, italic=italic, size=size, color=col or COL_TEXTE)
        _spacing(p, after=spa)

    _ctr("UNIVERSITÉ PARIS-SORBONNE", 14, bold=True, col=COL_BLEU_FONCE, spa=2)
    _ctr("Diplôme Universitaire — Sorbonne Data Analytics (DU SDA)",
         12, col=COL_BLEU_FONCE, spa=2)
    _ctr("Promotion 2025-2026", 11, italic=True, spa=18)
    _hrule(doc)
    doc.add_paragraph()
    _ctr("NOTE DE SYNTHÈSE" if is_note else "RAPPORT DE STAGE",
         20, bold=True, col=COL_BLEU_FONCE, spa=12)
    _ctr("Analyse Réseau et Prédiction de la Satisfaction Client",
         14, bold=True, col=COL_BLEU_MOYEN, spa=4)
    _ctr("dans l'Hôtellerie Indépendante", 13, italic=True,
         col=COL_BLEU_MOYEN, spa=4)
    _ctr("Application à l'Hôtel Aurore Paris Gare de Lyon",
         12, italic=True, spa=20)
    _hrule(doc)
    doc.add_paragraph()

    meta = [
        ("Stagiaire",           "[Votre Nom Prénom]"),
        ("Structure d'accueil", "Hôtel Aurore Paris Gare de Lyon — Paris 12ème"),
        ("Responsable stage",   "[Responsable à l'hôtel]"),
        ("Tuteur pédagogique",  "[Tuteur Université Paris-Sorbonne]"),
        ("Période",             "[Dates du stage]"),
        ("Date de remise",      "Avril 2026"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(meta):
        r0 = tbl.rows[i].cells[0]
        r0.text = ""
        _run(r0.paragraphs[0], lbl, bold=True, size=11, color=COL_BLEU_FONCE)
        _set_cell_bg(r0, "EEF2FF")
        _set_cell_border(r0)
        r1 = tbl.rows[i].cells[1]
        r1.text = ""
        _run(r1.paragraphs[0], val, size=11)
        _set_cell_border(r1)

    doc.add_paragraph()
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# Rendu tableau Markdown
# ══════════════════════════════════════════════════════════════════════════════

def _render_table(doc: Document, rows_raw: list):
    data = [r for r in rows_raw
            if not re.match(r"^\s*\|[-:\s|]+\|\s*$", r)]
    if not data:
        return

    def _parse(line):
        parts = [c.strip() for c in line.split("|")]
        return [p for p in parts if p != ""]

    header = _parse(data[0])
    ncols  = len(header)
    if ncols == 0:
        return

    tbl = doc.add_table(rows=len(data), cols=ncols)
    tbl.style = "Table Grid"

    for ri, raw in enumerate(data):
        cells_txt = _parse(raw)
        while len(cells_txt) < ncols:
            cells_txt.append("")
        for ci in range(ncols):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            raw_c = cells_txt[ci] if ci < len(cells_txt) else ""
            txt   = _strip(raw_c)
            is_bold = "**" in raw_c
            p = cell.paragraphs[0]
            _run(p, txt,
                 bold=(ri == 0 or is_bold),
                 size=10,
                 color=COL_BLANC if ri == 0 else COL_TEXTE)
            if ri == 0:
                _set_cell_bg(cell, "1F3C88")
            elif ri % 2 == 0:
                _set_cell_bg(cell, "EEF2FF")
            else:
                _set_cell_bg(cell, "FFFFFF")
            _set_cell_border(cell)

    p = doc.add_paragraph()
    _spacing(p, after=8)


# ══════════════════════════════════════════════════════════════════════════════
# Rendu bloc de code
# ══════════════════════════════════════════════════════════════════════════════

def _render_code(doc: Document, lines: list):
    if not lines:
        return
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    _set_cell_bg(cell, "F0F0F8")
    _set_cell_border(cell)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5A)
        _spacing(p, before=0, after=0)
    p = doc.add_paragraph()
    _spacing(p, after=6)


# ══════════════════════════════════════════════════════════════════════════════
# Moteur de conversion
# ══════════════════════════════════════════════════════════════════════════════

def _convert(md_path: Path, docx_path: Path, is_note: bool):
    print(f"  Lecture : {md_path.name}")
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.readlines()

    doc = Document()
    _configure(doc)
    _cover(doc, is_note)

    in_code   = False
    code_buf  = []
    in_table  = False
    table_buf = []
    skip      = True   # sauter l'en-tête de couverture MD

    def flush_code():
        nonlocal in_code, code_buf
        _render_code(doc, code_buf)
        code_buf = []
        in_code  = False

    def flush_table():
        nonlocal in_table, table_buf
        _render_table(doc, table_buf)
        table_buf = []
        in_table  = False

    for raw in lines:
        line = raw.rstrip("\n")

        # Sauter le bloc de couverture MD (re-généré en Word)
        if skip:
            if re.match(r"^#{1,4}\s+.+", line):
                skip = False
            else:
                continue

        # Blocs de code
        if line.strip().startswith("```"):
            if in_table:
                flush_table()
            if in_code:
                flush_code()
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # Tableaux
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table  = True
                table_buf = []
            table_buf.append(line)
            continue
        elif in_table:
            flush_table()

        # Ligne vide
        if not line.strip():
            p = doc.add_paragraph()
            _spacing(p, after=2)
            continue

        # Séparateurs ---
        if re.match(r"^-{3,}$", line.strip()):
            _hrule(doc)
            continue

        # Titres
        hm = re.match(r"^(#{1,5})\s+(.*)", line)
        if hm:
            level = len(hm.group(1))
            text  = _strip(hm.group(2))
            p = doc.add_paragraph()
            before = max(6, 20 - level * 3)
            after  = max(4, 14 - level * 2)
            _spacing(p, before=before, after=after)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _run(p, text, bold=True, size=18, color=COL_BLEU_FONCE)
                _hrule(doc)
            elif level == 2:
                _run(p, text, bold=True, size=14, color=COL_BLEU_FONCE)
            elif level == 3:
                _run(p, text, bold=True, size=12, color=COL_BLEU_MOYEN)
            elif level == 4:
                _run(p, text, bold=True, size=11, color=COL_BLEU_MOYEN)
            else:
                _run(p, text, bold=True, italic=True, size=11)
            continue

        # Citations
        if line.strip().startswith(">"):
            txt = _strip(line.strip().lstrip(">").strip())
            p   = doc.add_paragraph()
            _run(p, txt, italic=True, size=11,
                 color=RGBColor(0x44, 0x44, 0x88))
            p.paragraph_format.left_indent  = Cm(1.5)
            p.paragraph_format.right_indent = Cm(1.5)
            _spacing(p, before=4, after=4)
            continue

        # Listes
        bm = re.match(r"^(\s*)([-*+]|\d+[.)]) (.*)", line)
        if bm:
            indent = len(bm.group(1)) // 2
            prefix = f"{bm.group(2)} " if re.match(r"\d+[.)]", bm.group(2)) else "•  "
            txt    = _strip(bm.group(3))
            p = doc.add_paragraph()
            _run(p, prefix + txt, size=11)
            p.paragraph_format.left_indent       = Cm(1.0 + indent * 0.6)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            _spacing(p, before=1, after=2)
            continue

        # Paragraphe normal avec inline bold/italic/code
        txt = line.strip()
        if not txt:
            continue
        p = doc.add_paragraph()
        _spacing(p, before=0, after=4)
        for seg in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", txt):
            if seg.startswith("**") and seg.endswith("**"):
                _run(p, seg[2:-2], bold=True, size=11)
            elif seg.startswith("*") and seg.endswith("*"):
                _run(p, seg[1:-1], italic=True, size=11)
            elif seg.startswith("`") and seg.endswith("`"):
                _run(p, seg[1:-1], size=10, font="Courier New",
                     color=RGBColor(0x1A, 0x1A, 0x7A))
            else:
                seg = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", seg)
                if seg:
                    _run(p, seg, size=11)

    # Flush finaux
    if in_code:
        flush_code()
    if in_table:
        flush_table()

    doc.save(str(docx_path))
    size_kb = docx_path.stat().st_size // 1024
    print(f"  Sauvegardé : {docx_path.name}  ({size_kb} Ko)")
    return docx_path


# ══════════════════════════════════════════════════════════════════════════════
# Point d'entrée
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("  Génération Word — DU SDA 2025-2026")
    print("  Hôtel Aurore Paris Gare de Lyon")
    print("=" * 60)
    results = []
    for md, docx, note in [
        (RAPPORT_MD, RAPPORT_DOCX, False),
        (NOTE_MD,    NOTE_DOCX,    True),
    ]:
        if not md.exists():
            print(f"  ⚠  Introuvable : {md.name}")
            continue
        print()
        results.append(_convert(md, docx, is_note=note))

    print()
    print("─" * 60)
    for p in results:
        kb = Path(p).stat().st_size // 1024
        print(f"  ✅  {Path(p).name}  —  {kb} Ko")
    print("─" * 60)
    print("  Terminé.")


if __name__ == "__main__":
    main()

