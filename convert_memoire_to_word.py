#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversion du mémoire Markdown → Word (.docx)
Mémoire DU SDA 2025-2026 — Université Paris-Sorbonne
Hôtel Aurore Paris — Analyse Réseau & Satisfaction Client
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Chemins ──────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
MD_PATH  = BASE_DIR / "docs" / "MEMOIRE_MASTER_DATA_SCIENCE_2026.md"
OUT_PATH = BASE_DIR / "docs" / "MEMOIRE_MASTER_DATA_SCIENCE_2026.docx"

# ─── Couleurs ─────────────────────────────────────────────────────────────────
COLOR_TITLE   = RGBColor(0x1F, 0x39, 0x64)   # Bleu marine sorbonne
COLOR_H1      = RGBColor(0x1F, 0x39, 0x64)   # Bleu marine
COLOR_H2      = RGBColor(0x2E, 0x74, 0xB5)   # Bleu moyen
COLOR_H3      = RGBColor(0x2E, 0x74, 0xB5)   # Bleu moyen
COLOR_H4      = RGBColor(0x40, 0x40, 0x40)   # Gris foncé
COLOR_CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)   # Gris très clair


def set_cell_background(cell, hex_color: str):
    """Définit la couleur de fond d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document):
    """Ajoute une ligne horizontale (séparateur)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def configure_styles(doc: Document):
    """Configure les styles personnalisés du document."""
    styles = doc.styles

    # Style Normal de base
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(14)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_H1
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_H2
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_H3
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    # Heading 4
    h4 = styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = COLOR_H4
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    # Style "Code" pour blocs de code
    try:
        code_style = styles['Code']
    except KeyError:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.left_indent = Cm(0.5)

    # Style "Quote" pour citations
    try:
        quote_style = styles['Quote']
    except KeyError:
        quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = 'Calibri'
    quote_style.font.size = Pt(11)
    quote_style.font.italic = True
    quote_style.paragraph_format.left_indent = Cm(1.5)
    quote_style.paragraph_format.space_before = Pt(8)
    quote_style.paragraph_format.space_after = Pt(8)


def add_title_page(doc: Document):
    """Crée la page de titre stylisée."""
    # Espace en haut
    for _ in range(3):
        doc.add_paragraph()

    # Université
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Université Paris-Sorbonne")
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITLE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("DU SDA — Promotion 2025-2026")
    run2.font.name = 'Calibri'
    run2.font.size = Pt(13)
    run2.font.color.rgb = COLOR_H2

    add_horizontal_rule(doc)

    # Titre principal
    for _ in range(2):
        doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("ANALYSE RÉSEAU ET MODÉLISATION DE LA\nSATISFACTION CLIENT DANS LE SECTEUR HÔTELIER")
    run3.font.name = 'Calibri'
    run3.font.size = Pt(22)
    run3.font.bold = True
    run3.font.color.rgb = COLOR_TITLE

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Application à l'Hôtel Aurore Paris Gare de Lyon")
    run4.font.name = 'Calibri'
    run4.font.size = Pt(14)
    run4.font.italic = True
    run4.font.color.rgb = COLOR_H2

    add_horizontal_rule(doc)

    for _ in range(3):
        doc.add_paragraph()

    # Infos auteur — tableau
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    infos = [
        ("Auteur", "[Votre Nom Prénom]"),
        ("Directeur de mémoire", "[Nom du Directeur]"),
        ("Date de soutenance", "Juin 2026"),
        ("UFR", "Sciences Mathématiques et Informatique"),
        ("Spécialité", "Data Science & Intelligence Artificielle"),
    ]
    for i, (label, value) in enumerate(infos):
        row = table.rows[i]
        # Cellule label
        cell_label = row.cells[0]
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].bold = True
        cell_label.paragraphs[0].runs[0].font.name = 'Calibri'
        set_cell_background(cell_label, "DEEAF1")
        # Cellule valeur
        cell_value = row.cells[1]
        cell_value.text = value
        cell_value.paragraphs[0].runs[0].font.name = 'Calibri'

    doc.add_paragraph()

    # Citation
    quote_p = doc.add_paragraph()
    quote_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_p.paragraph_format.left_indent = Cm(2)
    quote_p.paragraph_format.right_indent = Cm(2)
    quote_run = quote_p.add_run(
        '« La valeur d\'un individu dans un réseau ne réside pas seulement dans ce qu\'il est,\n'
        'mais dans la façon dont il est connecté aux autres. »\n'
        '— Mark Newman, Networks: An Introduction, 2010'
    )
    quote_run.font.name = 'Calibri'
    quote_run.font.size = Pt(10)
    quote_run.font.italic = True
    quote_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Saut de page
    doc.add_page_break()


def apply_inline_formatting(run_or_para, text: str):
    """
    Applique le formatage inline (gras, italique) à un paragraphe.
    Retourne True si du formatage a été appliqué.
    """
    # Nettoyer les balises Markdown inline dans le texte brut
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    clean = re.sub(r'\*(.+?)\*', r'\1', clean)
    clean = re.sub(r'`(.+?)`', r'\1', clean)
    clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean)
    clean = re.sub(r'~~(.+?)~~', r'\1', clean)
    return clean


def add_formatted_paragraph(doc: Document, text: str, style: str = 'Normal',
                              bold: bool = False, italic: bool = False,
                              alignment=None, indent: float = 0) -> None:
    """Ajoute un paragraphe avec formatage inline Markdown basique."""
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Cm(indent)

    # Découper le texte selon les marqueurs Markdown
    pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[.+?\]\(.+?\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        run = p.add_run()
        run.font.name = 'Calibri'

        if part.startswith('***') and part.endswith('***'):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif re.match(r'\[(.+?)\]\(.+?\)', part):
            m = re.match(r'\[(.+?)\]\(.+?\)', part)
            run.text = m.group(1)
            run.underline = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            run.text = part

        if bold:
            run.bold = True
        if italic:
            run.italic = True


def add_table_from_markdown(doc: Document, lines: list):
    """Convertit un tableau Markdown en tableau Word."""
    # Filtrer les lignes de séparateur (|---|---|)
    data_lines = [l for l in lines if not re.match(r'^\s*\|[\s\-:]+\|', l)]
    if not data_lines:
        return

    rows_data = []
    for line in data_lines:
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        cells = [c.strip() for c in line.split('|')]
        rows_data.append(cells)

    if not rows_data:
        return

    max_cols = max(len(r) for r in rows_data)
    # Normaliser le nombre de colonnes
    rows_data = [r + [''] * (max_cols - len(r)) for r in rows_data]

    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            # Appliquer le formatage inline
            p = cell.paragraphs[0]
            pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
            parts = re.split(pattern, cell_text)
            for part in parts:
                if not part:
                    continue
                run = p.add_run()
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                if part.startswith('***') and part.endswith('***'):
                    run.text = part[3:-3]; run.bold = True; run.italic = True
                elif part.startswith('**') and part.endswith('**'):
                    run.text = part[2:-2]; run.bold = True
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    run.text = part[1:-1]; run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run.text = part[1:-1]; run.font.name = 'Courier New'
                else:
                    run.text = part

            # En-tête : fond coloré
            if i == 0:
                set_cell_background(cell, "2E74B5")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
            elif i % 2 == 0:
                set_cell_background(cell, "EBF3FB")

    doc.add_paragraph()


def parse_and_convert(doc: Document, md_content: str):
    """Parse le Markdown et remplit le document Word."""
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        raw_line = line
        stripped = line.strip()

        # ── Blocs de code ────────────────────────────────────────────────────
        if stripped.startswith('```') or stripped.startswith('````'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # Fin du bloc de code
                in_code_block = False
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(style='Code')
                    p.paragraph_format.left_indent = Cm(0.8)
                    # Fond gris clair via ombrage de paragraphe
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F2F2F2')
                    pPr.append(shd)
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8.5)
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # ── Sauts de page (--- seul sur une ligne) ───────────────────────────
        if stripped == '---' and len(stripped) == 3:
            # Vérifier si c'est vraiment un séparateur de section
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Tableaux Markdown ─────────────────────────────────────────────────
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                add_table_from_markdown(doc, table_lines)
                table_lines = []
            # Ne pas continuer — traiter la ligne courante

        # ── Titres ────────────────────────────────────────────────────────────
        h4_match = re.match(r'^####\s+(.*)', stripped)
        h3_match = re.match(r'^###\s+(.*)', stripped)
        h2_match = re.match(r'^##\s+(.*)', stripped)
        h1_match = re.match(r'^#\s+(.*)', stripped)

        if h1_match:
            title_text = apply_inline_formatting(None, h1_match.group(1))
            # Saut de page avant les chapitres principaux (chiffre seul)
            if re.match(r'^\d+\.', title_text):
                doc.add_page_break()
            p = doc.add_paragraph(style='Heading 1')
            run = p.add_run(title_text)
            run.font.color.rgb = COLOR_H1
            i += 1
            continue

        if h2_match:
            title_text = apply_inline_formatting(None, h2_match.group(1))
            p = doc.add_paragraph(style='Heading 2')
            run = p.add_run(title_text)
            run.font.color.rgb = COLOR_H2
            i += 1
            continue

        if h3_match:
            title_text = apply_inline_formatting(None, h3_match.group(1))
            p = doc.add_paragraph(style='Heading 3')
            run = p.add_run(title_text)
            run.font.color.rgb = COLOR_H3
            i += 1
            continue

        if h4_match:
            title_text = apply_inline_formatting(None, h4_match.group(1))
            p = doc.add_paragraph(style='Heading 4')
            run = p.add_run(title_text)
            i += 1
            continue

        # ── Lignes vides ──────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Citations (> texte) ───────────────────────────────────────────────
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            # Retirer le formatage Markdown
            quote_text = apply_inline_formatting(None, quote_text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
            i += 1
            continue

        # ── Listes ────────────────────────────────────────────────────────────
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            is_ordered = bool(re.match(r'\d+\.', list_match.group(2)))
            item_text = list_match.group(3)
            style = 'List Number' if is_ordered else 'List Bullet'
            add_formatted_paragraph(doc, item_text, style=style,
                                    indent=indent_level * 0.5)
            i += 1
            continue

        # ── Formules mathématiques ($$...$$) ──────────────────────────────────
        if stripped.startswith('$$') or re.match(r'^\$\$.+\$\$', stripped):
            formula_text = stripped.replace('$$', '').strip()
            if formula_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(formula_text)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(11)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
            i += 1
            continue

        # ── Paragraphe normal ─────────────────────────────────────────────────
        add_formatted_paragraph(doc, stripped)
        i += 1
        continue

    # Flush tableau si en fin de fichier
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)


def add_page_numbers(doc: Document):
    """Ajoute les numéros de page dans le pied de page."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()

    # Texte gauche
    run_left = p.add_run("Mémoire DU SDA 2025-2026 — Université Paris-Sorbonne")
    run_left.font.name = 'Calibri'
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    run_sep = p.add_run("   |   Page ")
    run_sep.font.name = 'Calibri'
    run_sep.font.size = Pt(8)
    run_sep.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Numéro de page automatique
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_num = p.add_run()
    run_num.font.size = Pt(8)
    run_num._r.append(fldChar1)
    run_num._r.append(instrText)
    run_num._r.append(fldChar2)


def set_page_margins(doc: Document):
    """Définit les marges de la page."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def main():
    print(f"📖 Lecture du mémoire : {MD_PATH}")
    md_content = MD_PATH.read_text(encoding='utf-8')

    print("📝 Création du document Word...")
    doc = Document()

    # Configuration
    set_page_margins(doc)
    configure_styles(doc)

    # Page de titre
    add_title_page(doc)

    # Contenu principal
    # On saute la première ligne (# MÉMOIRE DE FIN D'ÉTUDES) déjà dans la page de titre
    # et le bloc YAML de titre
    print("🔄 Conversion du contenu Markdown...")
    parse_and_convert(doc, md_content)

    # Pied de page
    add_page_numbers(doc)

    # Sauvegarde
    print(f"💾 Sauvegarde dans : {OUT_PATH}")
    doc.save(str(OUT_PATH))

    print(f"\n✅ Conversion réussie !")
    print(f"   Fichier créé : {OUT_PATH}")
    print(f"   Taille : {OUT_PATH.stat().st_size / 1024:.1f} Ko")


if __name__ == "__main__":
    main()

